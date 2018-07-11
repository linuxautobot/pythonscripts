import re
import csv
from docx import Document


def docx_replace_regex(doc_obj, regex , replace):

    for p in doc_obj.paragraphs:
        if regex.search(p.text):
            inline = p.runs
            # Loop added to work with runs (strings with same style)
            for i in range(len(inline)):
                if regex.search(inline[i].text):
                    text = regex.sub(replace, inline[i].text)
                    inline[i].text = text

    for table in doc_obj.tables:
        for row in table.rows:
            for cell in row.cells:
                docx_replace_regex(cell, regex , replace)



with open('a.csv') as csvfile:
    readCSV = csv.reader(csvfile)
    for row in readCSV:
         ab = row[0:2]
         replaced = " ".join(map(str, ab))
         print replaced
         regex1 = re.compile(r"xxx xxxx")
         replace1 = replaced
         filename = "test.docx"
         doc = Document(filename)
         docx_replace_regex(doc, regex1 , replace1)
         doc.save(replaced+'.docx')



